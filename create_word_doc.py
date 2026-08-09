import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Color Palette
    PRIMARY = RGBColor(15, 23, 42)    # Slate 900
    ACCENT = RGBColor(2, 132, 199)    # Sky 600
    TEAL = RGBColor(13, 148, 136)     # Teal 600
    DARK = RGBColor(30, 41, 59)       # Slate 800
    MUTED = RGBColor(100, 116, 139)   # Slate 500

    # Header / Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run("SMART SHIP PORT")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(10)
    run_sub = sub_p.add_run("3-Minute Presentation Script & Complete Project Walkthrough")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = ACCENT

    # Meta banner table
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_cell = meta_table.cell(0, 0)
    set_cell_background(meta_cell, "F0F9FF")
    set_cell_margins(meta_cell, top=140, bottom=140, left=180, right=180)
    meta_p = meta_cell.paragraphs[0]
    meta_p.paragraph_format.space_after = Pt(0)
    meta_p.add_run("System: ").bold = True
    meta_p.add_run("Real-Time AIS Vessel Tracker, Sonar Scanner HUD & Predictive Harbor Logistics\n")
    meta_p.add_run("Repository: ").bold = True
    meta_p.add_run("https://github.com/Sathiya2-coder/Smart-port | ")
    meta_p.add_run("Target Duration: ").bold = True
    meta_p.add_run("Exactly 3:00 Minutes")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section: Script
    sec1 = doc.add_paragraph()
    sec1.paragraph_format.space_before = Pt(14)
    sec1.paragraph_format.space_after = Pt(6)
    r_sec1 = sec1.add_run("1. Timed 3-Minute Presentation Script")
    r_sec1.font.name = "Arial"
    r_sec1.font.size = Pt(14)
    r_sec1.font.bold = True
    r_sec1.font.color.rgb = PRIMARY

    # Function to add timed presentation block
    def add_presentation_block(time_range, title, screen_action, speech_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        
        p_head = cell.paragraphs[0]
        p_head.paragraph_format.space_after = Pt(4)
        r_time = p_head.add_run(f"⏱️ {time_range} — {title}\n")
        r_time.font.name = "Arial"
        r_time.font.size = Pt(11)
        r_time.font.bold = True
        r_time.font.color.rgb = TEAL
        
        r_action = p_head.add_run(f"📺 Screen Action: {screen_action}\n")
        r_action.font.name = "Arial"
        r_action.font.size = Pt(9.5)
        r_action.font.italic = True
        r_action.font.color.rgb = MUTED

        p_speech = cell.add_paragraph()
        p_speech.paragraph_format.space_before = Pt(4)
        p_speech.paragraph_format.space_after = Pt(0)
        r_sp_label = p_speech.add_run("🎤 Spoken Dialogue:\n")
        r_sp_label.font.name = "Arial"
        r_sp_label.font.size = Pt(9.5)
        r_sp_label.font.bold = True
        r_sp_label.font.color.rgb = PRIMARY

        r_sp_text = p_speech.add_run(f'"{speech_text}"')
        r_sp_text.font.name = "Arial"
        r_sp_text.font.size = Pt(9.5)
        r_sp_text.font.color.rgb = DARK

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. Intro
    add_presentation_block(
        "0:00 – 0:30",
        "Introduction & The Problem / Home Page",
        "Open Home Page (Show hero title, live metrics bar, and system workflow)",
        "Hello everyone, today I am presenting Smart Ship Port — a next-generation Maritime Intelligence, Real-Time AIS Fleet Tracking, and Automated Berth Allocation platform.\n\n"
        "In global maritime shipping, ports face major bottlenecks: unpredictable vessel arrival delays, berth queue congestion, and fragmented tracking systems. Smart Ship Port solves this by combining real-time AIS telemetry streaming, hardware-accelerated marine cartography, tactical sonar scanning, and predictive AI decision models.\n\n"
        "On our Home Page, you can see live telemetry metrics tracking our active merchant fleet, real-time message throughput, and major Indian Ocean hubs, backed by an end-to-end 5-step data workflow from sensor ingestion to 1-click harbor gate execution."
    )

    # 2. Live Map
    add_presentation_block(
        "0:30 – 1:20",
        "Live Interactive AIS Map & Fleet Directory",
        "Click 'LAUNCH LIVE MAP' -> Click MV Thoothukudi Express -> Show glowing trajectory track",
        "Moving into our core Live AIS Navigation Map:\n\n"
        "On the left is our Fleet Directory, tracking merchant vessels including Container Ships, Oil Tankers, Bulk Carriers, and Offshore Vessels. Users can search by name, MMSI, or filter by operational status.\n\n"
        "On the right is our Marine Cartography Engine, powered by Leaflet and hardware-accelerated Canvas rendering. Notice that all vessels are actively sailing in real-time: our 60FPS motion engine computes live nautical velocity vectors, smoothly rotating vessel headings and gliding markers with GSAP interpolation.\n\n"
        "When I click any vessel — such as MV Thoothukudi Express — the camera executes a smooth fly-to, opening its complete vessel dossier and drawing its live glowing trajectory track line.\n\n"
        "We also monitor 10 Major Indian Ocean Ports — including JNPT Mumbai, Chennai, Cochin, Mundra, and Colombo — showing real-time occupancy percentages and congestion alerts."
    )

    # 3. Sonar HUD
    add_presentation_block(
        "1:20 – 2:05",
        "Tactical Sonar Radar Scanner HUD",
        "Click 'SONAR HUD' in top navbar -> Toggle range rings -> Click a target to navigate to map",
        "Next is one of our flagship innovations — the Tactical Sonar Radar HUD.\n\n"
        "Built with GSAP and HTML5 Canvas, it runs a continuous 360-degree acoustic radar beam sweep with real-time sinusoidal wave telemetry.\n\n"
        "Using true polar bearing trigonometry and Haversine geodesic distance formulas, the radar calculates exact nautical distance and bearing angles from the user's datum to every ship and port.\n\n"
        "Our 'ALL FLEET' auto-scaling mode fits the entire fleet across the circular radar disk, along with custom zoom ranges (1.5k km, 3k km, 5k km) and quick category filters for Ships, Ports, Moving, and Anchored vessels. Clicking any radar blip instantly closes the HUD and lands right on that ship's live navigation route on the map."
    )

    # 4. Predictive AI
    add_presentation_block(
        "2:05 – 2:40",
        "Predictive Logistics & AI Maritime Companion",
        "Click 'AI ANALYTICS' -> Show 6-hour forecast and AI chat advisory",
        "To optimize harbor turnaround, we integrated Predictive Logistics and an AI Maritime Companion powered by Google Gemini:\n\n"
        "1. Predictive Congestion Index: Forecasts berth delays, container dwell times, and yard stacking density 6 hours in advance.\n"
        "2. AI Maritime Companion: Features specialized personas for Port Managers, Navigators, and General Explorers. A port manager can ask for optimal berth assignments, and the AI computes spatial proximity, draft depth clearance, estimated fuel cost avoided, and gate queue reduction — providing explainable, actionable recommendations."
    )

    # 5. Conclusion
    add_presentation_block(
        "2:40 – 3:00",
        "Technology Stack Summary & Conclusion",
        "Show final dashboard view with sleek dark glassmorphic interface",
        "Under the hood, Smart Ship Port is engineered with React 18, Vite 5, Tailwind CSS, Leaflet Canvas, GSAP motion physics, WebSockets, and Geodesic spherical algorithms.\n\n"
        "Smart Ship Port transforms maritime logistics from reactive delays into an automated, predictive, zero-bottleneck ecosystem.\n\n"
        "Thank you! I am now happy to answer any questions."
    )

    # Section 2: Q&A Cheat Sheet Table
    sec2 = doc.add_paragraph()
    sec2.paragraph_format.space_before = Pt(14)
    sec2.paragraph_format.space_after = Pt(6)
    r_sec2 = sec2.add_run("2. Viva & Presentation Q&A Defense Matrix")
    r_sec2.font.name = "Arial"
    r_sec2.font.size = Pt(14)
    r_sec2.font.bold = True
    r_sec2.font.color.rgb = PRIMARY

    qa_data = [
        ("Examiner Question", "Strong Technical Answer"),
        ("What data source is used for vessel tracking?", "Live AIS WebSocket protocol (wss://stream.aisstream.io/v0/stream) with an automated zero-latency dead-reckoning fallback engine."),
        ("How do the ships move smoothly without lag?", "GSAP (GreenSock) coordinate vector interpolation over Course Over Ground (COG) and Speed Over Ground (SOG) running on requestAnimationFrame."),
        ("How does the Sonar HUD calculate distance and angles?", "Great-circle Haversine formula for spherical distance in km, and polar azimuth projection equations (x = r*sin θ, y = -r*cos θ)."),
        ("What makes the UI performant with 100+ vessels?", "Leaflet Hardware-Accelerated Canvas rendering (preferCanvas: true) and memory buffering that batches state updates 4 times/sec instead of re-rendering on every packet."),
        ("How is the AI assistant implemented?", "Google Gemini 2.5 Flash API with specialized maritime system instructions, geospatial distance prompts, and explainable cost-avoidance metrics.")
    ]

    qa_table = doc.add_table(rows=len(qa_data), cols=2)
    qa_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Inches(2.4), Inches(4.4)]
    for row_idx, (q, a) in enumerate(qa_data):
        row = qa_table.rows[row_idx]
        cell_q = row.cells[0]
        cell_a = row.cells[1]

        cell_q.width = col_widths[0]
        cell_a.width = col_widths[1]

        if row_idx == 0:
            set_cell_background(cell_q, "0F172A")
            set_cell_background(cell_a, "0F172A")
            pq = cell_q.paragraphs[0]
            pa = cell_a.paragraphs[0]
            rq = pq.add_run(q)
            ra = pa.add_run(a)
            rq.bold = True
            ra.bold = True
            rq.font.color.rgb = RGBColor(255, 255, 255)
            ra.font.color.rgb = RGBColor(255, 255, 255)
        else:
            bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell_q, bg)
            set_cell_background(cell_a, bg)
            pq = cell_q.paragraphs[0]
            pa = cell_a.paragraphs[0]
            rq = pq.add_run(q)
            ra = pa.add_run(a)
            rq.bold = True
            rq.font.color.rgb = PRIMARY
            ra.font.color.rgb = DARK

        set_cell_margins(cell_q, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_a, top=80, bottom=80, left=100, right=100)

    filename = "Smart_Ship_Port_Presentation_Script.docx"
    doc.save(filename)
    print("Word Document Successfully Generated:", filename)

if __name__ == "__main__":
    create_document()
